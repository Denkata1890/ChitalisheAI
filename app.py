import streamlit as st
import google.generativeai as genai
import os
from docx import Document
from io import BytesIO
import PyPDF2
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import io
from datetime import datetime


# 1. ФУНКЦИЯ ЗА СВЪРЗВАНЕ С DRIVE
def get_drive_service():
    info = st.secrets["gcp_service_account"]
    credentials = service_account.Credentials.from_service_account_info(info)
    scoped_credentials = credentials.with_scopes(['https://www.googleapis.com/auth/drive'])
    return build('drive', 'v3', credentials=scoped_credentials)


# 2. ФУНКЦИЯ ЗА СЪЗДАВАНЕ НА ПАПКА
def get_or_create_folder(folder_name, parent_folder_id=None):
    service = get_drive_service()
    query = f"name = '{folder_name}' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
    if parent_folder_id:
        query += f" and '{parent_folder_id}' in parents"

    results = service.files().list(q=query, fields="files(id, name)").execute()
    items = results.get('files', [])

    if not items:
        file_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        if parent_folder_id:
            file_metadata['parents'] = [parent_folder_id]
        folder = service.files().create(body=file_metadata, fields='id').execute()
        return folder.get('id')
    return items[0]['id']


# 3. ФУНКЦИЯ ЗА КАЧВАНЕ НА ФАЙЛ
def upload_to_drive(file_content, file_name, folder_id):
    service = get_drive_service()
    file_metadata = {
        'name': file_name,
        'parents': [folder_id]
    }
    media = MediaIoBaseUpload(io.BytesIO(file_content),
                              mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    service.files().create(body=file_metadata, media_body=media, fields='id').execute()


# 4. ФУНКЦИЯ ЗА КАЧВАНЕ НА PDF/TXT ФАЙЛОВЕ
def upload_source_file_to_drive(file_content, file_name, folder_id):
    service = get_drive_service()

    # Определяме MIME типа според разширението
    if file_name.endswith('.pdf'):
        mime_type = 'application/pdf'
    else:
        mime_type = 'text/plain'

    file_metadata = {
        'name': file_name,
        'parents': [folder_id]
    }
    media = MediaIoBaseUpload(io.BytesIO(file_content), mimetype=mime_type)
    service.files().create(body=file_metadata, media_body=media, fields='id').execute()


# 5. ПРОВЕРКА ЗА ПОТРЕБИТЕЛ
def check_password():
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False

    if st.session_state["authenticated"]:
        return True

    st.title("🔐 Вход в Chitalishe AI Pro")
    user = st.text_input("Потребителско име")
    password = st.text_input("Парола", type="password")

    if st.button("Влизане"):
        if user in USERS and USERS[user] == password:
            st.session_state["authenticated"] = True
            st.session_state["username"] = user
            st.rerun()
        else:
            st.error("❌ Грешно потребителско име или парола")
    return False


# 6. ПОМОЩНИ ФУНКЦИИ
def create_docx(text):
    doc = Document()
    doc.add_heading('Официален документ - Читалищен Секретар AI', 0)
    doc.add_paragraph(text)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


# 7. КОНФИГУРАЦИЯ
try:
    USERS = st.secrets["users"]
except:
    USERS = {"admin": "admin123"}

# Проверка за вход
if not check_password():
    st.stop()

# Информация за потребителя
if st.session_state["username"] == "admin":
    st.sidebar.warning("🛠️ Влезли сте като Администратор. Имате пълен достъп.")
else:
    st.sidebar.info(f"🏛️ Добре дошли, НЧ 'Светлина - 1861г.'")

# Създаваме основна папка за потребителя в Google Drive
try:
    # Първо създаваме основна папка "Chitalishe_AI_Pro"
    main_folder_id = get_or_create_folder("Chitalishe_AI_Pro")
    # След това създаваме потребителска подпапка
    user_folder_id = get_or_create_folder(f"user_{st.session_state['username']}", main_folder_id)
    # Създаваме подпапка за документи
    documents_folder_id = get_or_create_folder("Generated_Documents", user_folder_id)
    # Създаваме подпапка за архив
    archive_folder_id = get_or_create_folder("Archive", user_folder_id)
    st.sidebar.success("✅ Свързан с Google Drive")
except Exception as e:
    st.sidebar.error(f"❌ Грешка с Google Drive: {str(e)}")
    documents_folder_id = None
    archive_folder_id = None

st.sidebar.write(f"👤 Вписани сте като: **{st.session_state['username']}**")
if st.sidebar.button("Изход"):
    st.session_state["authenticated"] = False
    st.rerun()

# 8. КОНФИГУРАЦИЯ НА ИИ
API_KEY = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=API_KEY)

instruction = """Ти си професионален софтуерен асистент за управление на читалища. 
Твоята сила е, че комбинираш ОБЩИТЕ ЗАКОНИ (ЗНЧ) с ЛОКАЛНИЯ АРХИВ (Устав, заповеди, договори) на конкретното читалище.
Винаги отговаряй въз основа на предоставените документи. Ако липсва информация в локалния архив, провери в законите."""

model = genai.GenerativeModel('gemini-1.5-flash', system_instruction=instruction)


# 9. ЗАРЕЖДАНЕ НА ГЛОБАЛНАТА БАЗА
@st.cache_data
def get_global_knowledge():
    context = ""
    knowledge_dir = "ai_knowledge_base"
    if os.path.exists(knowledge_dir):
        for filename in sorted(os.listdir(knowledge_dir)):
            with open(os.path.join(knowledge_dir, filename), 'r', encoding='utf-8') as f:
                context += f"ИЗТОЧНИК {filename}:\n{f.read()}\n\n"
    return context


# 10. ИНТЕРФЕЙС
st.set_page_config(page_title="Chitalishe AI Pro", page_icon="🏛️", layout="wide")

# --- СТРАНИЧНА ЛЕНТА (УПРАВЛЕНИЕ НА АРХИВА) ---
st.sidebar.title("🔐 Личен архив на Читалището")
st.sidebar.write("Тук качвате документите, които са специфични само за вашата организация.")

uploaded_files = st.sidebar.file_uploader(
    "Качете Устав, договори или стари протоколи (PDF/TXT):",
    type=['pdf', 'txt'],
    accept_multiple_files=True
)

local_context = ""
if uploaded_files:
    st.sidebar.success(f"Заредени са {len(uploaded_files)} лични документа.")

    # Записваме качените файлове в Google Drive
    if archive_folder_id:
        for uploaded_file in uploaded_files:
            try:
                # Прочитаме съдържанието
                file_content = uploaded_file.getvalue()
                # Качваме в Drive
                upload_source_file_to_drive(file_content, uploaded_file.name, archive_folder_id)
                st.sidebar.info(f"📁 {uploaded_file.name} записан в Drive")
            except Exception as e:
                st.sidebar.warning(f"⚠️ Неуспешен запис на {uploaded_file.name} в Drive: {str(e)}")

    # Обработваме съдържанието за текущата сесия
    for uploaded_file in uploaded_files:
        if uploaded_file.name.endswith('.pdf'):
            local_context += f"\nДОКУМЕНТ ОТ ВАШИЯ АРХИВ ({uploaded_file.name}):\n" + extract_text_from_pdf(
                uploaded_file)
        else:
            local_context += f"\nДОКУМЕНТ ОТ ВАШИЯ АРХИВ ({uploaded_file.name}):\n" + str(uploaded_file.read(), "utf-8")

# --- ИЗБОР НА РЕЖИМ ---
st.sidebar.divider()
mode = st.sidebar.radio("Действие:", ["⚖️ Консултация и Търсене", "📝 Създаване на протокол (Аудио/Текст)"])

# --- ОСНОВЕН ЕКРАН ---
st.title("🏛️ Читалищен Секретар AI - Професионална версия")

if mode == "⚖️ Консултация и Търсене":
    st.info("ИИ анализира едновременно законите и вашите качени документи.")
    user_input = st.chat_input("Напр. 'Какво казва нашият Устав за избор на председател?'")

    if user_input:
        with st.chat_message("user"):
            st.markdown(user_input)
        with st.chat_message("assistant"):
            with st.spinner("Проверявам в архивите..."):
                global_kb = get_global_knowledge()
                full_prompt = f"ГЛОБАЛНИ ЗАКОНИ:\n{global_kb}\n\nЛОКАЛЕН АРХИВ НА ЧИТАЛИЩЕТО:\n{local_context}\n\nВЪПРОС: {user_input}"
                response = model.generate_content(full_prompt)
                st.markdown(response.text)

                # Създаваме Word документ
                docx = create_docx(response.text)

                # Бутон за изтегляне
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        "📥 Изтегли като Word",
                        data=docx,
                        file_name=f"consultation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    )

                # Запис в Google Drive
                if documents_folder_id:
                    with col2:
                        if st.button("💾 Запиши в Google Drive"):
                            try:
                                docx.seek(0)
                                file_name = f"consultation_{st.session_state['username']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                                upload_to_drive(docx.getvalue(), file_name, documents_folder_id)
                                st.success(f"✅ Документът е записан в Google Drive като '{file_name}'")
                            except Exception as e:
                                st.error(f"❌ Грешка при запис в Drive: {str(e)}")

else:
    # РЕЖИМ ПРОТОКОЛИСТ
    st.subheader("Генератор на протоколи")
    audio_file = st.file_uploader("Качете запис:", type=['mp3', 'wav', 'm4a'])
    text_input = st.text_area("Или поставете текст тук:", height=150)

    if st.button("Генерирай официален документ"):
        with st.chat_message("assistant"):
            with st.spinner("Работя по документа..."):
                input_data = f"Използвай този личен архив за контекст:\n{local_context}\n\nНаправи протокол от това събрание:\n{text_input if text_input else 'Аудио файл (необработен)'}"

                if audio_file:
                    # Записваме аудиото в Drive ако има
                    if archive_folder_id:
                        audio_content = audio_file.getvalue()
                        upload_source_file_to_drive(audio_content,
                                                    f"audio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.mp3",
                                                    archive_folder_id)

                response = model.generate_content(input_data)
                st.markdown(response.text)

                # Създаваме Word документ
                docx = create_docx(response.text)

                # Бутон за изтегляне
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        "📥 Изтегли готовия файл",
                        data=docx,
                        file_name=f"protocol_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    )

                # Запис в Google Drive
                if documents_folder_id:
                    with col2:
                        if st.button("💾 Запиши протокола в Google Drive"):
                            try:
                                docx.seek(0)
                                file_name = f"protocol_{st.session_state['username']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                                upload_to_drive(docx.getvalue(), file_name, documents_folder_id)
                                st.success(f"✅ Протоколът е записан в Google Drive като '{file_name}'")
                            except Exception as e:
                                st.error(f"❌ Грешка при запис в Drive: {str(e)}")